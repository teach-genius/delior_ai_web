# ── Stdlib ────────────────────────────────────────────────────────────────────
import asyncio
import json
import logging
import os
import re
import shutil
import tempfile
import textwrap
import unicodedata
from pathlib import Path

# ── Third-party ───────────────────────────────────────────────────────────────
import requests
from docx import Document
from pdf2image import convert_from_path
from PIL import Image, ImageDraw, ImageFont
from rapidfuzz import fuzz, process

# ── Django ────────────────────────────────────────────────────────────────────
from django.conf import settings
from django.template.loader import render_to_string
from django.utils.timezone import now

# ── Internes projet ───────────────────────────────────────────────────────────
from django.db import IntegrityError
# ─────────────────────────────────────────────────────────────────────────────
logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# CONSTANTES
# ─────────────────────────────────────────────────────────────────────────────
SIMILARITY_THRESHOLD = 82

# ─────────────────────────────────────────────────────────────────────────────
# PROMPTS LLM
# ─────────────────────────────────────────────────────────────────────────────
QUERY_TEMPLATE = """
Analyse le CV par rapport à la fiche de poste.
CONTEXTE : {contexte}
OBJECTIF : évaluer cohérence candidat-poste, identifier points forts, écarts, risques, produire analyse structurée.
FICHE DE POSTE :
{job_description}

CV :
{cv_text}

INSTRUCTIONS :
- Se baser uniquement sur le CV
- Comparer compétences, expérience et parcours
- Identifier adéquations et écarts
- Signaler toute information manquante
- Pas de supposition
- Produire le rapport structuré selon le format JSON attendu
"""
# ─────────────────────────────────────────────────────────────────────────────
# HELPERS GÉNÉRIQUES
# ─────────────────────────────────────────────────────────────────────────────
def _safe_str(value) -> str:
    return value.strip() if isinstance(value, str) else ""

def _safe_list(value) -> list:
    return value if isinstance(value, list) else []

# ─────────────────────────────────────────────────────────────────────────────
# NORMALISATION — VILLE
# ─────────────────────────────────────────────────────────────────────────────
def normalize_ville(ville: str, timeout: int = 3) -> str:
    if not ville or not isinstance(ville, str):
        return ""
    ville = ville.strip()
    try:
        ville = ville.encode("latin-1").decode("utf-8")
    except (UnicodeEncodeError, UnicodeDecodeError):
        pass
    ville = unicodedata.normalize("NFC", ville).title()
    try:
        response = requests.get(
            "https://nominatim.openstreetmap.org/search",
            params={
                "q":              ville,
                "format":         "json",
                "addressdetails": 1,
                "limit":          1,
                "featuretype":    "city",
            },
            headers={
                "User-Agent":      "cv-app/1.0",  
                "Accept-Language": "fr",     
            },
            timeout=timeout,
        )
        response.raise_for_status()
        results = response.json()

        if not results:
            logger.debug("normalize_ville : aucun résultat Nominatim pour %r", ville)
            return ville

        address      = results[0].get("address", {})
        nom_officiel = (
            address.get("city")
            or address.get("town")
            or address.get("village")
            or address.get("county")
            or ville
        )
        nom_officiel = unicodedata.normalize("NFC", nom_officiel).strip()
        logger.debug("normalize_ville : %r → %r", ville, nom_officiel)
        return nom_officiel

    except requests.exceptions.Timeout:
        logger.warning("normalize_ville : timeout Nominatim pour %r", ville)
        return ville
    except Exception as exc:
        logger.warning("normalize_ville : erreur API pour %r — %s", ville, exc)
        return ville

def normalize_competence(nom: str, existing: list[str]) -> str:
    if not nom or not isinstance(nom, str):
        return ""

    nom = unicodedata.normalize("NFC", nom.lower().strip())

    if not nom:
        return ""

    if nom in existing:
        return nom

    if existing:
        match, score, _ = process.extractOne(
            nom,
            existing,
            scorer=fuzz.token_sort_ratio,
        )
        if score >= SIMILARITY_THRESHOLD:
            logger.debug(
                "normalize_competence fuzzy : %r → %r (score=%d)", nom, match, score
            )
            return match
    return nom.lower().strip()


# ─────────────────────────────────────────────────────────────────────────────
# ANALYSE LLM
# ────────────────────────────────────────────────────────────────────────────
async def analyse_cv_async(
    cv_text: str,
    job_description: str
) -> str:
    from aiagent.processing.chaine_analyse_cv import extract_analyse_cv
    """Version asynchrone de ``_run_analyse`` via ``asyncio.to_thread``."""
    return await asyncio.to_thread(extract_analyse_cv, cv_text, job_description)


# ─────────────────────────────────────────────────────────────────────────────
# RAPPORT JSON
# ─────────────────────────────────────────────────────────────────────────────

def generate_repport(cv: dict, query: dict) -> str | None:
    from aiagent.processing.chaine_analyse_cv import extract_analyse_cv
    from aiagent.recommender.recommender_sys import (
    cv_json_to_text_for_reasoning_impl,
    offre_json_to_text_for_reasoning_impl
    )

    if not isinstance(cv, dict) or not query:
    	raise ValueError("cv (dict) et query (dict) sont obligatoires.")

    cv_text = cv_json_to_text_for_reasoning_impl(cv)
    offre_text = offre_json_to_text_for_reasoning_impl(query)

    json_str = extract_analyse_cv(cv_text,offre_text)

    if not json_str:
        logger.warning("generate_repport : aucun JSON extrait. Réponse brute : %r", json_str)
        return None
    # Validation minimale des clés attendues
    try:
        parsed = json_str if isinstance(json_str, dict) else json.loads(json_str)
        if "candidat" not in parsed or "recommandation_finale" not in parsed:
            logger.warning(
                "generate_repport : JSON incomplet — clés reçues : %s",
                list(parsed.keys()),
            )
    except json.JSONDecodeError as e:
        logger.error(
            "generate_repport : JSON malformé — %s | extrait : %r", e, json_str[:300]
        )
        return None

    return json_str

# ─────────────────────────────────────────────────────────────────────────────
# CONVERSIONS VERS PDF
# ─────────────────────────────────────────────────────────────────────────────

def _image_to_pdf(input_path: str, output_path: str) -> bool:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas

        A4_W, A4_H = A4  # points (595 x 842)

        with Image.open(input_path) as img:
            img     = img.convert("RGB")
            img_w, img_h = img.size

            # Calcul du ratio pour tenir dans A4 avec marges de 20pt
            max_w, max_h = A4_W - 40, A4_H - 40
            ratio        = min(max_w / img_w, max_h / img_h)
            draw_w       = img_w * ratio
            draw_h       = img_h * ratio
            x            = (A4_W - draw_w) / 2
            y            = (A4_H - draw_h) / 2

            # Sauvegarde temporaire en JPEG (requis par reportlab drawImage)
            tmp_jpg = output_path + "_tmp.jpg"
            img.save(tmp_jpg, "JPEG", quality=90)

        c = canvas.Canvas(output_path, pagesize=A4)
        c.drawImage(tmp_jpg, x, y, width=draw_w, height=draw_h)
        c.save()

        os.remove(tmp_jpg)
        logger.debug("_image_to_pdf : %s → %s", input_path, output_path)
        return True

    except Exception as exc:
        logger.exception("_image_to_pdf : erreur — %s", exc)
        return False


def _docx_to_pdf_pure(input_path: str, output_path: str) -> bool:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

        doc    = Document(input_path)
        styles = getSampleStyleSheet()

        # Styles personnalisés
        style_normal = ParagraphStyle(
            "cvNormal",
            parent     = styles["Normal"],
            fontSize   = 10,
            leading    = 14,
            spaceAfter = 2,
        )
        style_h1 = ParagraphStyle(
            "cvH1",
            parent     = styles["Heading1"],
            fontSize   = 14,
            leading    = 18,
            spaceAfter = 6,
            textColor  = (0.1, 0.1, 0.4),
        )
        style_h2 = ParagraphStyle(
            "cvH2",
            parent     = styles["Heading2"],
            fontSize   = 11,
            leading    = 15,
            spaceAfter = 4,
        )

        STYLE_MAP = {
            "Title":     style_h1,
            "Heading 1": style_h1,
            "Heading 2": style_h2,
            "Heading 3": style_h2,
        }

        story = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                story.append(Spacer(1, 0.15 * cm))
                continue

            style = STYLE_MAP.get(para.style.name, style_normal)

            # Échappement des caractères spéciaux XML pour reportlab
            safe_text = (
                text
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            story.append(Paragraph(safe_text, style))

        pdf_doc = SimpleDocTemplate(
            output_path,
            pagesize     = A4,
            leftMargin   = 2 * cm,
            rightMargin  = 2 * cm,
            topMargin    = 2 * cm,
            bottomMargin = 2 * cm,
        )
        pdf_doc.build(story)

        logger.debug("_docx_to_pdf_pure : %s → %s", input_path, output_path)
        return True

    except Exception as exc:
        logger.exception("_docx_to_pdf_pure : erreur — %s", exc)
        return False


def convert_to_pdf(file_path: str) -> str | None:
    if not file_path or not os.path.isfile(file_path):
        logger.warning("convert_to_pdf : fichier introuvable — %s", file_path)
        return None

    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return file_path   # déjà un PDF, rien à faire

    if ext == ".doc":
        logger.warning("convert_to_pdf : format .doc non supporté sans LibreOffice.")
        return None

    tmp_dir  = tempfile.mkdtemp(prefix="cv_pdf_")
    out_path = os.path.join(tmp_dir, Path(file_path).stem + ".pdf")

    if ext == ".docx":
        success = _docx_to_pdf_pure(file_path, out_path)
    elif ext in (".png", ".jpg", ".jpeg", ".webp"):
        success = _image_to_pdf(file_path, out_path)
    else:
        logger.warning("convert_to_pdf : format non supporté — %s", ext)
        shutil.rmtree(tmp_dir, ignore_errors=True)
        return None

    if not success:
        shutil.rmtree(tmp_dir, ignore_errors=True)
        return None

    logger.debug("convert_to_pdf : %s → %s", file_path, out_path)
    return out_path


# ─────────────────────────────────────────────────────────────────────────────
# PDF / DOCX / IMAGE PREVIEW
# ─────────────────────────────────────────────────────────────────────────────

def _docx_to_preview_image(input_path: str, output_path: str) -> bool:
    """
    Génère une image JPEG de prévisualisation d'un DOCX via python-docx + Pillow.
    Simule la première "page" en rendant le texte extrait sur un fond blanc A4.

    Args:
        input_path:  Chemin absolu vers le fichier DOCX.
        output_path: Chemin absolu du JPEG de sortie.

    Returns:
        True si succès, False sinon.
    """
    try:
        doc = Document(input_path)

        # Extraction des premiers paragraphes non vides
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append((text, para.style.name))
            if len(lines) >= 40:   # ~40 lignes visibles max
                break

        if not lines:
            logger.warning("_docx_to_preview_image : document vide — %s", input_path)
            return False

        # ── Rendu image A4 ────────────────────────────────────────────────────
        W, H       = 794, 1123     # A4 à 96 dpi
        MARGIN     = 60
        bg_color   = (255, 255, 255)
        text_color = (30, 30, 30)
        head_color = (10, 10, 80)

        img  = Image.new("RGB", (W, H), color=bg_color)
        draw = ImageDraw.Draw(img)

        # Fontes système DejaVu (fallback : fonte PIL par défaut)
        try:
            font_normal = ImageFont.truetype(
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 13
            )
            font_bold = ImageFont.truetype(
                "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 15
            )
        except IOError:
            font_normal = ImageFont.load_default()
            font_bold   = font_normal

        y = MARGIN
        for text, style in lines:
            if y > H - MARGIN:
                break

            is_heading = "Heading" in style or "Title" in style
            font       = font_bold if is_heading else font_normal
            color      = head_color if is_heading else text_color
            size       = 17 if "Title" in style else (15 if is_heading else 13)

            # Recharge la fonte à la bonne taille si TTF disponible
            try:
                font = ImageFont.truetype(font.path, size)
            except Exception:
                pass

            max_chars = (W - 2 * MARGIN) // (size // 2 + 2)
            wrapped   = textwrap.wrap(text, width=max(20, max_chars))

            for line in wrapped:
                if y > H - MARGIN:
                    break
                draw.text((MARGIN, y), line, fill=color, font=font)
                y += size + 4

            y += 4  # espacement inter-paragraphe

        # Bordure légère
        draw.rectangle([10, 10, W - 10, H - 10], outline=(200, 200, 200), width=1)

        img.save(output_path, "JPEG", quality=85)
        return True

    except Exception as exc:
        logger.exception("_docx_to_preview_image : erreur — %s", exc)
        return False


def pdf_preview(path_file: str, output_folder: str = "cv_previews") -> str | None:
    if not path_file or not os.path.isfile(path_file):
        logger.warning("pdf_preview : fichier introuvable — %s", path_file)
        return None

    full_output_folder = os.path.join(settings.MEDIA_ROOT, output_folder)
    os.makedirs(full_output_folder, exist_ok=True)

    ext              = Path(path_file).suffix.lower()
    preview_filename = Path(path_file).stem + ".jpg"
    preview_path     = os.path.join(full_output_folder, preview_filename)
    relative_path    = os.path.join(output_folder, preview_filename)

    # ── PNG / JPG / WEBP → conversion directe via Pillow ─────────────────────
    if ext in (".png", ".jpg", ".jpeg", ".webp"):
        try:
            with Image.open(path_file) as img:
                img = img.convert("RGB")     # supprime le canal alpha éventuel
                img.thumbnail((794, 1123))   # redimensionne proportionnellement (A4 max)
                img.save(preview_path, "JPEG", quality=85)
            return relative_path
        except Exception as exc:
            logger.exception("pdf_preview : erreur image — %s", exc)
            return None

    # ── DOCX → rendu texte via python-docx + Pillow ───────────────────────────
    if ext == ".docx":
        success = _docx_to_preview_image(path_file, preview_path)
        return relative_path if success else None

    # ── DOC → non supporté sans LibreOffice ──────────────────────────────────
    if ext == ".doc":
        logger.warning("pdf_preview : format .doc non supporté sans LibreOffice.")
        return None

    # ── PDF → première page via pdf2image ─────────────────────────────────────
    try:
        pages = convert_from_path(path_file, dpi=150, first_page=1, last_page=1)
    except Exception as exc:
        logger.exception("pdf_preview : erreur convert_from_path — %s", exc)
        return None

    if not pages:
        logger.warning("pdf_preview : aucune page extraite de %s", path_file)
        return None

    try:
        pages[0].save(preview_path, "JPEG", quality=85)
    except Exception as exc:
        logger.exception("pdf_preview : erreur sauvegarde JPEG — %s", exc)
        return None

    return relative_path

# ─────────────────────────────────────────────────────────────────────────────
# RAPPORT HTML
# ─────────────────────────────────────────────────────────────────────────────
def generate_html_report_tool(context: dict, template_name: str) -> str:
    if not context or not isinstance(context, dict):
        raise ValueError("Le contexte fourni est invalide ou vide.")
    if not template_name or not isinstance(template_name, str):
        raise ValueError("template_name est obligatoire.")

    context.setdefault("meta", {})
    context["meta"].setdefault("date_analyse", now().isoformat())

    return render_to_string(template_name, context)



from aiagent.processing.extractor import cv_to_text
# ─────────────────────────────────────────────────────────────────────────────
# EXTRACTION / STRUCTURATION CV
# ─────────────────────────────────────────────────────────────────────────────
async def get_text_brut_cv_candidat(path_cv: str) -> str:
    from aiagent.processing.extractor import cv_to_text
    if not path_cv or not os.path.isfile(path_cv):
        raise FileNotFoundError(f"Fichier CV introuvable : {path_cv}")
    return await cv_to_text(path=path_cv)


def get_cv_text_brut_to_json(text_cv: str) -> dict | None:
    from aiagent.processing.chaine_extraction_cv import extract_cv_complet
    if not text_cv or not text_cv.strip():
        logger.warning("get_cv_text_brut_to_json : texte vide reçu.")
        return None
    return extract_cv_complet(text_cv)

# ─────────────────────────────────────────────────────────────────────────────
# INFOS CANDIDAT
# ─────────────────────────────────────────────────────────────────────────────
def get_infos_json_candidat(cv_json: dict) -> dict:
    if not isinstance(cv_json, dict):
        logger.warning(
            "get_infos_json_candidat : cv_json invalide (%s)", type(cv_json)
        )
        return {}

    identite            = cv_json.get("identite") or {}
    nom_complet         = _safe_str(identite.get("nom_complet"))
    titre_professionnel = _safe_str(identite.get("titre_professionnel"))

    contact   = identite.get("contact") or {}
    ville     = _safe_str(contact.get("ville"))
    pays      = _safe_str(contact.get("pays"))
    telephone = _safe_str(contact.get("telephone"))
    email     = _safe_str(contact.get("email"))
    liens     = _safe_list(contact.get("liens"))

    classification   = cv_json.get("classification") or {}
    domaine          = _safe_str(classification.get("domaine"))
    secteur          = _safe_str(classification.get("secteur"))
    niveau           = _safe_str(classification.get("niveau"))

    contrat_souhaite = _safe_str(cv_json.get("contrat_souhaite"))
    profil_resume    = _safe_str(cv_json.get("profil_resume"))

    formation = [
        {
            "diplome":       _safe_str(f.get("diplome")),
            "etablissement": _safe_str(f.get("etablissement")),
            "periode":       _safe_str(f.get("periode")),
            "lieu":          _safe_str(f.get("lieu")),
        }
        for f in _safe_list(cv_json.get("formation"))
        if isinstance(f, dict)
    ]

    experience = [
        {
            "poste":      _safe_str(e.get("poste")),
            "entreprise": _safe_str(e.get("entreprise")),
            "periode":    _safe_str(e.get("periode")),
            "missions":   _safe_list(e.get("missions")),
        }
        for e in _safe_list(cv_json.get("experience_professionnelle"))
        if isinstance(e, dict)
    ]

    competences  = cv_json.get("competences") or {}
    savoir_faire = [
        c.strip()
        for c in _safe_list(competences.get("savoir_faire"))
        if isinstance(c, str) and c.strip()
    ]
    savoir_etre = [
        c.strip()
        for c in _safe_list(competences.get("savoir_etre"))
        if isinstance(c, str) and c.strip()
    ]
    langues = [
        {"langue": _safe_str(l.get("langue")), "niveau": _safe_str(l.get("niveau"))}
        for l in _safe_list(competences.get("langues"))
        if isinstance(l, dict) and l.get("langue")
    ]

    pac = cv_json.get("projets_et_certifications") or {}
    if not isinstance(pac, dict):
        pac = {}

    projets_raw        = _safe_list(pac.get("projets"))        or _safe_list(cv_json.get("projets"))
    certifications_raw = _safe_list(pac.get("certifications")) or _safe_list(cv_json.get("certifications"))

    projets = [
        {
            "nom":          _safe_str(p.get("nom")),
            "description":  _safe_str(p.get("description")),
            "technologies": _safe_str(p.get("technologies")),
        }
        for p in projets_raw
        if isinstance(p, dict)
    ]

    certifications = [
        {
            "nom":       _safe_str(c.get("nom")),
            "organisme": _safe_str(c.get("organisme")),
            "annee":     _safe_str(c.get("annee")),
        }
        for c in certifications_raw
        if isinstance(c, dict)
    ]

    return {
        "nom_complet":         nom_complet,
        "titre_professionnel": titre_professionnel,
        "ville":               ville,
        "pays":                pays,
        "telephone":           telephone,
        "email":               email,
        "liens":               liens,
        "domaine":             domaine,
        "secteur":             secteur,
        "niveau":              niveau,
        "contrat_souhaite":    contrat_souhaite,
        "profil_resume":       profil_resume,
        "formation":           formation,
        "experience":          experience,
        "savoir_faire":        savoir_faire,
        "savoir_etre":         savoir_etre,
        "langues":             langues,
        "projets":             projets,
        "certifications":      certifications,
    }


def _safe_str_lower(value) -> str | None:
    if not value:
        return None 
    return str(value).lower().strip() or None

# ─────────────────────────────────────────────────────────────────────────────
# SAUVEGARDE CANDIDAT
# ─────────────────────────────────────────────────────────────────────────────
def save_candidat_from_cv_path(file_path: str, source_file=None):
    from asgiref.sync import async_to_sync
    from django.core.files.base import ContentFile
    from django.db import transaction, IntegrityError
    from candidat.models import CandidatCV, Competence

    if not file_path or not os.path.isfile(file_path):
        raise FileNotFoundError(f"Fichier introuvable : {file_path}")

    # ── 1. Conversion vers PDF ────────────────────────────────────────────────
    pdf_path = convert_to_pdf(file_path)
    if not pdf_path:
        raise ValueError(f"Format non supporté ou conversion PDF impossible : {file_path}")

    pdf_is_tmp = pdf_path != file_path

    try:
        # ── 2. Extraction texte ───────────────────────────────────────────────
        text_brut = async_to_sync(get_text_brut_cv_candidat)(pdf_path)
        if not text_brut or not text_brut.strip():
            raise ValueError("Impossible d'extraire le texte du CV.")

        # ── 3. Structuration LLM ──────────────────────────────────────────────
        cv_struct = get_cv_text_brut_to_json(text_brut)
        if not cv_struct:
            raise ValueError("Impossible de structurer le CV.")

        # ── 4. Extraction des données candidat ────────────────────────────────
        data = get_infos_json_candidat(cv_json=cv_struct)
        logger.debug("Données candidat extraites : %s", data)

        email     = data.get("email", "").lower().strip() or None
        telephone = data.get("telephone", "").lower().strip() or None

        if not email and not telephone:
            raise ValueError("Ni email ni téléphone trouvé dans le CV — candidat ignoré.")

        lookup = {"email": email} if email else {"telephone": telephone}

        # ── 5. get_or_create avec gestion race condition Celery ───────────────
        try:
            with transaction.atomic():
                candidat, created = CandidatCV.objects.get_or_create(
                    **lookup,
                    defaults={
                        "nom_complet":          _safe_str_lower(data.get("nom_complet")),
                        "email":                _safe_str_lower(data.get("email")),
                        "telephone":            _safe_str_lower(data.get("telephone")),
                        "ville":                _safe_str_lower(data.get("ville")),
                        "pays":                 _safe_str_lower(data.get("pays")),
                        "titre_professionnel":  _safe_str_lower(data.get("titre_professionnel")),
                        "donnees_structurees":  data,
                        "domaine":              _safe_str_lower(data.get("domaine")),
                        "secteur":              _safe_str_lower(data.get("secteur")),
                        "niveau":               _safe_str_lower(data.get("niveau")),
                        "contrat_souhaite":     _safe_str_lower(data.get("contrat_souhaite")),
                        "resume_profil":        _safe_str_lower(data.get("profil_resume")),
                    },
                )
        except IntegrityError:
            # Race condition entre workers Celery — transaction saine ici
            candidat = CandidatCV.objects.get(**lookup)
            created  = False

        # ── 6. Mise à jour si candidat existant ───────────────────────────────
        if not created:
            candidat.nom_complet          = _safe_str_lower(data.get("nom_complet"))
            candidat.email                = _safe_str_lower(data.get("email"))
            candidat.telephone            = _safe_str_lower(data.get("telephone"))
            candidat.ville                = _safe_str_lower(data.get("ville"))
            candidat.pays                 = _safe_str_lower(data.get("pays", ""))
            candidat.titre_professionnel  = _safe_str_lower(data.get("titre_professionnel"))
            candidat.donnees_structurees  = data
            candidat.domaine              = _safe_str_lower(data.get("domaine"))
            candidat.secteur              = _safe_str_lower(data.get("secteur"))
            candidat.niveau               = _safe_str_lower(data.get("niveau"))
            candidat.resume_profil        = _safe_str_lower(data.get("profil_resume"))
            for field in (candidat.fichier_pdf_origine, candidat.preview_image):
                try:
                    if field and field.name and os.path.exists(field.path):
                        os.remove(field.path)
                except Exception as exc:
                    logger.warning("Impossible de supprimer l'ancien fichier : %s", exc)

        # ── 7. Compétences + PDF + Preview dans un seul atomic ────────────────
        with transaction.atomic():

            # Compétences normalisées
            candidat.competences.clear()
            existing_competences = list(Competence.objects.values_list("nom", flat=True))

            for nom_comp in data.get("savoir_faire", []):
                if not isinstance(nom_comp, str) or not nom_comp.strip():
                    continue
                nom_norm = normalize_competence(nom_comp, existing=existing_competences)
                if not nom_norm:
                    continue
                comp, comp_created = Competence.objects.get_or_create(nom=nom_norm)
                if comp_created:
                    existing_competences.append(nom_norm)
                candidat.competences.add(comp)

            # Sauvegarde PDF
            safe_name = f"{candidat.candidat_id}.pdf"
            try:
                if source_file and not pdf_is_tmp:
                    candidat.fichier_pdf_origine.save(safe_name, source_file, save=False)
                else:
                    with open(pdf_path, "rb") as f:
                        candidat.fichier_pdf_origine.save(
                            safe_name, ContentFile(f.read()), save=False
                        )
            except Exception as exc:
                raise ValueError(f"Erreur lors de la sauvegarde du PDF : {exc}") from exc

            # Preview JPEG
            preview_path = pdf_preview(pdf_path)
            if preview_path:
                candidat.preview_image = preview_path
            else:
                logger.warning(
                    "save_candidat_from_cv_path : preview non générée pour %s", file_path
                )

            candidat.save()

    finally:
        # Nettoyage du dossier temporaire de conversion
        if pdf_is_tmp:
            try:
                shutil.rmtree(Path(pdf_path).parent, ignore_errors=True)
                logger.debug(
                    "save_candidat_from_cv_path : dossier tmp supprimé — %s",
                    Path(pdf_path).parent,
                )
            except Exception as exc:
                logger.warning(
                    "save_candidat_from_cv_path : impossible de supprimer le tmp — %s", exc
                )

    return candidat, created
